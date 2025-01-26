import datetime
import tkinter as tk
from tkinter import messagebox
from tkinter import END
from tkinter import ttk
from ctypes import windll
import docx
from docx.shared import Cm
import tkcalendar
from os import getcwd

windll.shcore.SetProcessDpiAwareness(1)


# formating date
def date_format(date):
    for i, w in enumerate(date):
        date[i] = w.strip("0")

    if int(date[0]) < 10:
        tmp = str(date[0])
        date[0] = "0" + tmp

    if int(date[1]) < 10:
        tmp = str(date[1])
        date[1] = "0" + tmp

    tmp = str(date[2])
    date[2] = "20" + tmp

    return f"{date[1]}.{date[0]}.{date[2]}"


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        
        # colors
        self.bg_color = "#2D4356"
        self.widget_color = "#A76F6F"
        self.text_color = "#EAB2A0"
        self.insert_text_color = "#20262E"

        # calculation app placement on screen
        s_w = self.winfo_screenwidth()
        s_h = self.winfo_screenheight()
        x = round((s_w / 2) - 700)
        y = round((s_h / 2) - 350)

        # window parameters
        self.geometry("1400x700" + "+" + str(x) + "+" + str(y))
        self.title("WYCENA")
        self.config(background=self.bg_color)
        self.resizable(False, False)
        
        # variable initialization
        self.all_data = []
        self.date = None
        self.editing = False
        self.curr_index = None
        
        # widgets initialization
        label_entry_name = tk.Label(self,
                                    text="Nazwa",
                                    background=self.bg_color,
                                    foreground=self.text_color,
                                    font=("Consolas", 20)
                                    )

        self.entry_name = tk.Entry(self,
                                   width=51,
                                   background=self.widget_color,
                                   foreground=self.insert_text_color,
                                   cursor="plus",
                                   insertbackground=self.insert_text_color,
                                   font=("Consolas", 22)
                                   )

        label_entry_amount = tk.Label(self,
                                      text="Ilość",
                                      background=self.bg_color,
                                      foreground=self.text_color,
                                      font=("Consolas", 20)
                                      )

        self.entry_amount = tk.Entry(self,
                                     width=5,
                                     background=self.widget_color,
                                     foreground=self.insert_text_color,
                                     cursor="plus",
                                     insertbackground=self.insert_text_color,
                                     font=("Consolas", 22)
                                     )

        label_entry_price = tk.Label(self,
                                     text="Cena",
                                     background=self.bg_color,
                                     foreground=self.text_color,
                                     font=("Consolas", 20)
                                     )

        self.entry_price = tk.Entry(self,
                                    width=7,
                                    background=self.widget_color,
                                    foreground=self.insert_text_color,
                                    cursor="plus",
                                    insertbackground=self.insert_text_color,
                                    font=("Consolas", 22)
                                    )

        button_make_word_file = tk.Button(self,
                                          text="STWÓRZ PLIK",
                                          background=self.text_color,
                                          foreground=self.bg_color,
                                          activebackground=self.bg_color,
                                          activeforeground=self.text_color,
                                          font=("Consolas", 20),
                                          command=self.print_docx
                                          )

        self.listbox_of_all_data = tk.Listbox(self,
                                              width=92,
                                              height=10,
                                              background=self.bg_color,
                                              foreground=self.text_color,
                                              font=("Consolas", 20),
                                              selectbackground=self.text_color,
                                              selectforeground=self.bg_color
                                              )

        self.button_delete = tk.Button(self,
                                       text="❌",
                                       background=self.text_color,
                                       foreground=self.bg_color,
                                       activebackground=self.bg_color,
                                       activeforeground="#f00",
                                       font=("Consolas", 12),
                                       width=3,
                                       borderwidth=0,
                                       command=self.delete
                                       )

        button_edit = tk.Button(self,
                                text="🖊",
                                background=self.text_color,
                                foreground=self.bg_color,
                                activebackground=self.bg_color,
                                activeforeground="#00f",
                                font=("Consolas", 12),
                                width=3,
                                borderwidth=0,
                                command=self.edit
                                )

        self.button_down = tk.Button(self,
                                     text="⬇",
                                     background=self.text_color,
                                     foreground=self.bg_color,
                                     activebackground=self.bg_color,
                                     activeforeground="#0f0",
                                     font=("Consolas", 12),
                                     width=3,
                                     borderwidth=0,
                                     command=self.down
                                     )

        self.button_up = tk.Button(self,
                                   text="⬆",
                                   background=self.text_color,
                                   foreground=self.bg_color,
                                   activebackground=self.bg_color,
                                   activeforeground="#0f0",
                                   font=("Consolas", 12),
                                   width=3,
                                   borderwidth=0,
                                   command=self.up
                                   )

        button_date = tk.Button(self,
                                text="🕓",
                                background=self.text_color,
                                foreground=self.bg_color,
                                activebackground=self.bg_color,
                                activeforeground="yellow",
                                font=("Consolas", 12),
                                width=3,
                                borderwidth=0,
                                command=self.toplevel_show
                                )

        button_info = tk.Button(self,
                                text="ℹ",
                                background=self.bg_color,
                                foreground=self.text_color,
                                activebackground=self.bg_color,
                                activeforeground="#00f",
                                font=("Consolas", 15),
                                width=3,
                                borderwidth=0,
                                command=self.toplevel_info_show
                                )
        
        # action binding
        self.entry_name.bind("<Return>", self.data_validation)
        self.entry_amount.bind("<Return>", self.data_validation)
        self.entry_price.bind("<Return>", self.data_validation)
        
        # layout
        label_entry_name.place(x=30, y=10)
        self.entry_name.place(x=30, y=50)
        label_entry_amount.place(x=1090, y=10)
        self.entry_amount.place(x=1090, y=50)
        label_entry_price.place(x=1230, y=10)
        self.entry_price.place(x=1230, y=50)

        button_make_word_file.place(x=600, y=150)

        self.listbox_of_all_data.place(x=5, y=300)

        self.button_delete.place(x=1350, y=250)
        button_edit.place(x=1300, y=250)
        self.button_down.place(x=1250, y=250)
        self.button_up.place(x=1200, y=250)
        button_date.place(x=10, y=250)
        button_info.place(x=1360, y=0)

        # toplevel (calendar) displaying
        self.toplevel = tk.Toplevel(self)
        self.toplevel.title("DATA WAŻNOŚCI")
        
        tl_x = round((s_w / 2) - 200)
        tl_y = round((s_h / 2) - 250)

        self.toplevel.geometry("400x500" + "+" + str(tl_x) + "+" + str(tl_y))
        self.toplevel.resizable(False, False)
        self.toplevel.config(background="#7512e6")
        self.toplevel.overrideredirect(True)
        
        # toplevel widgets initialization
        tl_label_title = tk.Label(self.toplevel,
                                  text="ZAZNACZ DATĘ WAŻNOŚCI:",
                                  background="#7512e6",
                                  foreground="#233443",
                                  font=("Consolas", 18)
                                  )

        self.tl_calendar = tkcalendar.Calendar(self.toplevel,
                                               selectmode="day",
                                               year=datetime.date.today().year,
                                               month=datetime.date.today().month,
                                               day=datetime.date.today().day,
                                               background=self.bg_color,
                                               foreground=self.text_color,
                                               selectbackgorund=self.bg_color,
                                               normalbackground="#7512e6",
                                               weekendbackground="#540ca6",
                                               weekendforeground=self.widget_color,
                                               selectbackground=self.text_color,
                                               selectforeground=self.bg_color
                                               )

        tl_button_submit_time = tk.Button(self.toplevel,
                                          text="ZATWIERDŹ",
                                          background=self.text_color,
                                          foreground=self.bg_color,
                                          activebackground=self.bg_color,
                                          activeforeground="yellow",
                                          font=("Consolas", 12),
                                          command=self.toplevel_hide
                                          )
        # layout
        tl_label_title.pack(side="top", pady=15)
        self.tl_calendar.pack(side="top", padx=15, fill="both", expand=True)
        tl_button_submit_time.pack(side="bottom", pady=10)

        # hiding toplevel
        self.toplevel.withdraw()
        
        # toplevel (information) displaying
        self.toplevel_info = tk.Toplevel(self)
        self.toplevel_info.title("DATA WAŻNOŚCI")

        tl_x = round((s_w / 2) - 200)
        tl_y = round((s_h / 2) - 250)

        self.toplevel_info.geometry("400x500" + "+" + str(tl_x) + "+" + str(tl_y))
        self.toplevel_info.resizable(False, False)
        self.toplevel_info.config(background="#7512e6")
        self.toplevel_info.overrideredirect(True)
        
        # toplevel widgets initialization
        tli_sep = ttk.Separator(self.toplevel_info,
                                orient="horizontal"
                                )

        tli_button_close = tk.Button(self.toplevel_info,
                                     text="❌",
                                     background=self.bg_color,
                                     foreground=self.text_color,
                                     activebackground=self.bg_color,
                                     activeforeground="#f00",
                                     font=("Consolas", 12),
                                     borderwidth=0,
                                     command=self.toplevel_info.withdraw
                                     )

        tli_sep_close = ttk.Separator(self.toplevel_info,
                                      orient="horizontal"
                                      )

        tli_button_date = tk.Button(self.toplevel_info,
                                    text="🕓",
                                    background=self.text_color,
                                    foreground=self.bg_color,
                                    activebackground=self.bg_color,
                                    activeforeground="yellow",
                                    font=("Consolas", 12),
                                    borderwidth=0,
                                    width=3
                                    )

        tli_label_date = tk.Label(self.toplevel_info,
                                  text="- po naciśnięciu przycisku pojawi\n się kalendarz do wybrania\n daty ważności "
                                       "oferty",
                                  background="#7512e6",
                                  foreground=self.text_color,
                                  font=("Consolas", 10)
                                  )

        tli_sep_date = ttk.Separator(self.toplevel_info,
                                     orient="horizontal"
                                     )

        tli_button_up = tk.Button(self.toplevel_info,
                                  text="⬆",
                                  background=self.text_color,
                                  foreground=self.bg_color,
                                  activebackground=self.bg_color,
                                  activeforeground="#0f0",
                                  font=("Consolas", 12),
                                  borderwidth=0,
                                  width=3
                                  )

        tli_label_up = tk.Label(self.toplevel_info,
                                text="- po zaznaczeniu wpisu i naciśnięcia\n przycisku wpis przesunie się do góry",
                                background="#7512e6",
                                foreground=self.text_color,
                                font=("Consolas", 10)
                                )

        tli_sep_up = ttk.Separator(self.toplevel_info,
                                   orient="horizontal"
                                   )

        tli_button_down = tk.Button(self.toplevel_info,
                                    text="⬇",
                                    background=self.text_color,
                                    foreground=self.bg_color,
                                    activebackground=self.bg_color,
                                    activeforeground="#0f0",
                                    font=("Consolas", 12),
                                    borderwidth=0,
                                    width=3
                                    )

        tli_label_down = tk.Label(self.toplevel_info,
                                  text="- po zaznaczeniu wpisu i naciśnięcia\n przycisku wpis przesunie się w dół",
                                  background="#7512e6",
                                  foreground=self.text_color,
                                  font=("Consolas", 10)
                                  )

        tli_sep_down = ttk.Separator(self.toplevel_info,
                                     orient="horizontal"
                                     )

        tli_button_edit = tk.Button(self.toplevel_info,
                                    text="🖊",
                                    background=self.text_color,
                                    foreground=self.bg_color,
                                    activebackground=self.bg_color,
                                    activeforeground="#00f",
                                    font=("Consolas", 12),
                                    borderwidth=0,
                                    width=3
                                    )

        tli_label_edit = tk.Label(self.toplevel_info,
                                  text="- po zaznaczeniu wpisu i naciśnięcia\n przycisku wpis pojawi się w polach\n do "
                                       "wpisywania, a po dokonaniu\n modyfikacji zmiany można\n zatwierdzić ENTER-em",
                                  background="#7512e6",
                                  foreground=self.text_color,
                                  font=("Consolas", 10)
                                  )

        tli_sep_edit = ttk.Separator(self.toplevel_info,
                                     orient="horizontal"
                                     )

        tli_button_delete = tk.Button(self.toplevel_info,
                                      text="❌",
                                      background=self.text_color,
                                      foreground=self.bg_color,
                                      activebackground=self.bg_color,
                                      activeforeground="#f00",
                                      font=("Consolas", 12),
                                      borderwidth=0,
                                      width=3
                                      )

        tli_label_delete = tk.Label(self.toplevel_info,
                                    text="- po zaznaczeniu wpisu i naciśnięcia\n wpis zostanie usunięty",
                                    background="#7512e6",
                                    foreground=self.text_color,
                                    font=("Consolas", 10)
                                    )

        tli_sep_delete = ttk.Separator(self.toplevel_info,
                                       orient="horizontal"
                                       )
        
        # layout
        tli_sep.place(x=0, y=0, width=400)
        tli_button_close.place(x=360, y=0)
        tli_sep_close.place(x=0, y=35, width=400)

        tli_button_date.place(x=10, y=70)
        tli_label_date.place(x=60, y=70)
        tli_sep_date.place(x=0, y=140, width=400)

        tli_button_up.place(x=10, y=150)
        tli_label_up.place(x=60, y=150)
        tli_sep_up.place(x=0, y=210, width=400)

        tli_button_down.place(x=10, y=220)
        tli_label_down.place(x=60, y=220)
        tli_sep_down.place(x=0, y=280, width=400)

        tli_button_edit.place(x=10, y=290)
        tli_label_edit.place(x=60, y=290)
        tli_sep_edit.place(x=0, y=400, width=400)

        tli_button_delete.place(x=10, y=410)
        tli_label_delete.place(x=60, y=410)
        tli_sep_delete.place(x=0, y=498, width=400)

        # hiding toplevel
        self.toplevel_info.withdraw()
        self.deiconify()
    
    # validates data from entries
    def data_validation(self, *args, **kwargs):
        try:
            name = str(self.entry_name.get())
            amount = int(self.entry_amount.get().strip())
            price = float(self.entry_price.get().strip().replace(",", "."))
            if name == '':
                messagebox.showerror("PUSTE WARTOŚCI", "NIE WPROWADZONO WARTOŚCI")
            else:
                self.clear_entry()
                self.add(name, amount, price)

        except ValueError:
            messagebox.showerror("BŁĄD WARTOŚCI", "WPROWADZONO BŁĘDNĄ WARTOŚĆ")

        except:
            messagebox.showerror("BŁĄD", "NIEZNANY BŁĄD")

    # clears entries
    def clear_entry(self):
        self.entry_name.delete(0, END)
        self.entry_amount.delete(0, END)
        self.entry_price.delete(0, END)

    # adds product to listbox
    def add(self, name, amount, price):
        if not self.editing:
            arr = [name, amount, price]
            self.all_data.append(arr)
            self.listbox_of_all_data.insert(END, str(arr[0]) + "     " + str(arr[1]) + "     " + str(arr[2]))
        else:
            self.button_down.config(command=self.down)
            self.button_up.config(command=self.up)
            self.button_delete.config(command=self.delete)

            self.all_data[self.curr_index][0] = name
            self.all_data[self.curr_index][1] = amount
            self.all_data[self.curr_index][2] = price
            self.listbox_of_all_data.insert(self.curr_index, str(self.all_data[self.curr_index][0]) + "     "
                                            + str(self.all_data[self.curr_index][1])
                                            + "     " + str(self.all_data[self.curr_index][2]))
            self.listbox_of_all_data.delete(self.curr_index + 1)
        self.editing = False
    
    # shows toplevel (calendar)
    def toplevel_show(self, *args, **kwargs):
        self.toplevel.deiconify()

    # hides toplevel (calendar)
    def toplevel_hide(self, *args, **kwargs):
        self.toplevel.withdraw()
        self.submit_date()

    # shows toplevel (information)
    def toplevel_info_show(self, *args, **kwargs):
        self.toplevel_info.deiconify()
    
    # date getting
    def submit_date(self, *args, **kwargs):
        self.date = self.tl_calendar.get_date()
        self.date = self.date.split("/")

    # deletes chosen product from listbox
    def delete(self, *args, **kwargs):
        index = self.listbox_of_all_data.curselection()
        if index == ():
            messagebox.showinfo("INFO", "NIE ZAZNACZONO WPISU")
        else:
            del self.all_data[index[0]]
            self.listbox_of_all_data.delete(index[0])

    # moves product up on listbox
    def up(self, *args, **kwargs):
        index = self.listbox_of_all_data.curselection()

        if index == ():
            messagebox.showinfo("INFO", "NIE ZAZNACZONO WPISU")
        else:
            if index[0] != 0 and index != ():
                index = int(index[0])

                self.all_data[index], self.all_data[index - 1] = self.all_data[index - 1], self.all_data[index]

                self.listbox_of_all_data.insert(index, str(self.all_data[index][0]) + "     "
                                                + str(self.all_data[index][1])
                                                + "     " + str(self.all_data[index][2]))
                self.listbox_of_all_data.insert(index - 1, str(self.all_data[index - 1][0]) + "     "
                                                + str(self.all_data[index - 1][1])
                                                + "     " + str(self.all_data[index - 1][2]))

                self.listbox_of_all_data.delete(index + 1)
                self.listbox_of_all_data.delete(index + 1)

            else:
                messagebox.showinfo("INFO", "WPIS JEST JUŻ NA SAMEJ GÓRZE")
    
    # moves product down on listbox
    def down(self, *args, **kwargs):
        index = self.listbox_of_all_data.curselection()

        if index == ():
            messagebox.showinfo("INFO", "NIE ZAZNACZONO WPISU")
        else:
            if index[0] != len(self.all_data) - 1 and index != ():
                index = int(index[0])

                self.all_data[index], self.all_data[index + 1] = self.all_data[index + 1], self.all_data[index]

                self.listbox_of_all_data.insert(index, str(self.all_data[index][0]) + "     "
                                                + str(self.all_data[index][1])
                                                + "     " + str(self.all_data[index][2]))
                self.listbox_of_all_data.insert(index + 1, str(self.all_data[index + 1][0]) + "     "
                                                + str(self.all_data[index + 1][1])
                                                + "     " + str(self.all_data[index + 1][2]))

                self.listbox_of_all_data.delete(index + 2)
                self.listbox_of_all_data.delete(index + 2)

            else:
                messagebox.showinfo("INFO", "WPIS JEST JUŻ NA SAMYM DOLE")
    
    # edits selected on listbox product
    def edit(self, *args, **kwargs):
        self.editing = True
        self.curr_index = self.listbox_of_all_data.curselection()
        if self.curr_index == ():
            messagebox.showinfo("INFO", "NIE ZAZNACZONO WPISU")
        else:
            self.button_down.config(command="")
            self.button_up.config(command="")
            self.button_delete.config(command="")

            self.curr_index = self.curr_index[0]
            self.entry_name.delete(0, END)
            self.entry_amount.delete(0, END)
            self.entry_price.delete(0, END)

            self.entry_name.insert(0, self.all_data[self.curr_index][0])
            self.entry_amount.insert(0, self.all_data[self.curr_index][1])
            self.entry_price.insert(0, self.all_data[self.curr_index][2])
    
    # creates .docx file
    def print_docx(self, *args, **kwargs):

        if self.all_data:
            if self.date is not None:
                doc = docx.Document(getcwd()+"\\template.docx")
                table = doc.tables[0]
                for i in range(len(self.all_data)):
                    table.add_row()

                for i in range(len(self.all_data)):
                    row = table.rows[i + 1].cells
                    row[0].text = str(self.all_data[i][0])
                    row[2].text = str(self.all_data[i][1])
                    row[1].text = str(self.all_data[i][2]) + "zł"
                    row[3].text = str(round(self.all_data[i][1] * self.all_data[i][2], 2)) + "zł"

                for row in table.rows:
                    row.height = Cm(2.09)

                curr_date = datetime.datetime.now().strftime("%x")
                curr_date = curr_date.split("/")

                cd = date_format(curr_date)
                ed = date_format(self.date)

                rep = {
                    "%curr_date%": cd,
                    "%exp_date%": ed
                }

                for par in doc.paragraphs:
                    for key in rep:
                        par.style = "ST2"
                        par.text = par.text.replace(key, rep[key])

                for i in range(len(self.all_data)):
                    sum_all = round(self.all_data[i][1] * self.all_data[i][2], 2)
                    perc_sum_all = round(sum_all + sum_all * 0.23, 2)
                    p = doc.add_paragraph(f"{self.all_data[i][0]} w ilości {self.all_data[i][1]} sztuk:\n")
                    p.style = "ST"
                    r = p.add_run(f"{sum_all} zł netto (+23% VAT) - {perc_sum_all} zł brutto\n")
                    r.bold = True

                messagebox.showinfo("PLIK", "PLIK ZOSTAŁ STWORZONY")
                doc.save(f"wycena {cd}.docx")
            else:
                messagebox.showinfo("INFO", "PODAJ DATĘ")
        else:
            messagebox.showinfo("INFO", "NIE MA JESZCZE ŻADNYCH WPISÓW")


if __name__ == '__main__':
    app = App()
    app.mainloop()
